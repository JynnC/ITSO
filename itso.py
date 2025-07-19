from flask import Flask, render_template, jsonify, redirect, url_for, request, session, flash, make_response
import mysql.connector 
from flask import flash, redirect, url_for
from mysql.connector import Error
from werkzeug.security import generate_password_hash, check_password_hash
from flask import jsonify, request, send_file
from datetime import datetime, timedelta
import pandas as pd
import matplotlib.pyplot as plt
import io
from matplotlib.backends.backend_pdf import PdfPages
from PyPDF2 import PdfReader, PdfWriter
import weasyprint
import os
import calendar
from flask import session
from werkzeug.utils import secure_filename
from flask import send_from_directory
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import secrets
import string
import logging
import sys
from werkzeug.security import generate_password_hash
from flask import current_app
import traceback
from io import BytesIO
from flask import send_file
from docx import Document
from docx.shared import Inches
from flask_mail import Mail, Message


app = Flask(__name__)
app.secret_key = 'itso_group'

app.config['MAIL_SERVER'] = 'smtp.gmail.com' 
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'casilijynn@gmail.com'  
app.config['MAIL_PASSWORD'] = 'mkbu ugek npvo troy'  
app.config['MAIL_DEFAULT_SENDER'] = 'casilijynn@gmail.com'

mail = Mail(app)

UPLOAD_FOLDER = 'static/uploads/profile_pictures'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def get_db_connection():
    try:
        return mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="ITSO_RMS"
        )
    except Error as e:
        print("Database Connection Error: ", e)
        return None


@app.route('/')
@app.route('/login', methods=['GET', 'POST'])
def login():
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return render_template('login.html')
    
    cursor = db.cursor(dictionary=True)

    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()

        if not email or not password:
            flash('Please enter both email and password', 'error')
            return render_template('login.html')

        cursor.execute("""
            SELECT id, password, is_temp_password, role, is_verified 
            FROM users 
            WHERE email = %s
        """, (email,))
        user = cursor.fetchone()

        if user and check_password_hash(user['password'], password):
            session['user_id'] = user['id']
            session['role'] = user['role']

            if user['role'] == 'Coordinator' and not user['is_verified']:
                flash('Your account is pending verification.', 'warning')
                return render_template('login.html')

            return redirect(url_for('dashboard'))
        else:
            flash('Invalid email or password', 'error')

    cursor.close()
    db.close()
    return render_template('login.html')

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    """
    Handles the forgot password request.
    GET: Renders the forgot password form.
    POST: Processes the email submission and sends reset link.
    """
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        
        if not email:
            flash('Please enter your email address', 'error')
            return render_template('login.html', show_forgot_modal=True)
        
        db = get_db_connection()
        if not db:
            flash('Database Connection Failed', 'error')
            return render_template('login.html', show_forgot_modal=True)
        
        cursor = db.cursor(dictionary=True)
        
        cursor.execute("SELECT id FROM users WHERE email = %s", (email,))
        user = cursor.fetchone()
        
        if not user:
            flash('If your email is registered, you will receive password reset instructions shortly.', 'success')
            cursor.close()
            db.close()
            return render_template('login.html')
        

        try:
            print("About to calculate token expiry")
            token = secrets.token_urlsafe(32)
            token_expiry = datetime.now() + timedelta(hours=1)
            print(f"Token expiry calculated: {token_expiry}")
            
            cursor.execute("""
                INSERT INTO password_reset_tokens (user_id, token, expiry_date)
                VALUES (%s, %s, %s)
            """, (user['id'], token, token_expiry))
            db.commit()
        except Exception as e:
            db.rollback()
            print(f"Error in forgot_password: {str(e)}")
            import traceback
            traceback.print_exc()  # This will print the full traceback
            flash('An error occurred. Please try again later.', 'error')
            cursor.close()
            db.close()
            return render_template('login.html', show_forgot_modal=True)
        
        reset_link = url_for('reset_password', token=token, _external=True)
        send_password_reset_email(email, reset_link)
        
        flash('If your email is registered, you will receive password reset instructions shortly.', 'success')
        cursor.close()
        db.close()
        return render_template('login.html')
    
    return render_template('login.html', show_forgot_modal=True)


@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    """
    Handles the password reset process.
    GET: Validates token and shows reset form.
    POST: Updates the user's password.
    """
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    cursor = db.cursor(dictionary=True)
    
    cursor.execute("""
        SELECT user_id, expiry_date FROM password_reset_tokens
        WHERE token = %s AND used = 0
    """, (token,))
    token_info = cursor.fetchone()
    
    if not token_info or token_info['expiry_date'] < datetime.now():
        flash('Invalid or expired password reset link. Please request a new one.', 'error')
        cursor.close()
        db.close()
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        new_password = request.form.get('new_password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()
        
        if not new_password or len(new_password) < 8:
            flash('Password must be at least 8 characters long', 'error')
            return render_template('reset_password.html', token=token)
        
        if new_password != confirm_password:
            flash('Passwords do not match', 'error')
            return render_template('reset_password.html', token=token)
        
        hashed_password = generate_password_hash(new_password)
        
        try:
            cursor.execute("""
                UPDATE users SET password = %s WHERE id = %s
            """, (hashed_password, token_info['user_id']))
            
            cursor.execute("""
                UPDATE password_reset_tokens SET used = 1 WHERE token = %s
            """, (token,))
            
            db.commit()
            flash('Your password has been updated successfully. Please log in with your new password.', 'success')
            cursor.close()
            db.close()
            return redirect(url_for('login'))
        except Exception as e:
            db.rollback()
            flash('An error occurred. Please try again.', 'error')
            cursor.close()
            db.close()
            return render_template('reset_password.html', token=token)
    
    cursor.close()
    db.close()
    return render_template('reset_password.html', token=token)


def send_password_reset_email(email, reset_link):
    """
    Sends a password reset email to the user.
    
    Args:
        email (str): The recipient's email address
        reset_link (str): The password reset link
    """
    try:
        subject = "Password Reset Request"
        body = f"""
        You requested a password reset for your account.
        
        Please click the following link to reset your password:
        {reset_link}
        
        This link is valid for 1 hour.
        
        If you did not request a password reset, please ignore this email.
        """
        
        msg = Message(
            subject=subject,
            recipients=[email],
            body=body
        )
        
        mail.send(msg)
        print(f"Reset link for {email}: {reset_link}")
        return True
    except Exception as e:
        print(f"Error sending email: {str(e)}")
        return False

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        role = request.form.get('role')
        fullname = request.form.get('fullname')
        email = request.form.get('email')
        department = request.form.get('department')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm-password')

        if not all([role, fullname, email, department, password, confirm_password]):
            flash('All fields are required', 'error')
            return redirect(url_for('register'))

        if password != confirm_password:
            flash('Passwords do not match', 'error')
            return redirect(url_for('register'))

        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("SELECT * FROM users WHERE email = %s", (email,))
        if cursor.fetchone():
            flash('Email already registered', 'error')
            cursor.close()
            db.close()
            return redirect(url_for('register'))

        hashed_password = generate_password_hash(password)
        is_verified = (role.lower() == 'admin')  # Admin is automatically verified
        role = role.capitalize()

        try:
            cursor.execute("""
                INSERT INTO users (full_name, email, role, password, is_verified, department)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (fullname, email, role, hashed_password, is_verified, department))
            user_id = cursor.lastrowid

            cursor.execute("""
                INSERT INTO user_profiles (user_id, first_name, last_name, email, department)
                VALUES (%s, %s, %s, %s, %s)
            """, (user_id, fullname.split(' ')[0], ' '.join(fullname.split(' ')[1:]), email, department))

            db.commit()

            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))

        except Error as e:
            db.rollback()
            flash(f'Registration failed: {str(e)}', 'error')

        finally:
            cursor.close()
            db.close()

    return render_template('register.html')



@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out successfully', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']

        cursor.execute("""
            SELECT u.*, up.department as user_department 
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return redirect(url_for('login'))

        full_name = user['full_name'] if user else 'User'
        is_admin = user['role'] == 'Admin'
        user_department = user.get('department') or user.get('user_department')

        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        profile['profile_picture_url'] = url_for('static', filename=f'uploads/profile_pictures/{profile["profile_picture"]}') if profile.get('profile_picture') else url_for('static', filename='avatar.png')

        
        if is_admin:
            cursor.execute("""
                SELECT category, COUNT(*) as total_copies
                FROM document_metadata
                GROUP BY category
            """)
        else:
            cursor.execute("""
                SELECT dm.category, COUNT(*) as total_copies
                FROM document_metadata dm
                JOIN documents d ON dm.document_id = d.id
                WHERE d.department = %s
                GROUP BY dm.category
            """, (user_department,))
            
        data = cursor.fetchall()
        document_counts = {row['category']: row['total_copies'] for row in data}

        if is_admin:
            cursor.execute("""
                SELECT al.activity, al.timestamp, u.full_name
                FROM activity_logs al
                JOIN users u ON al.user_id = u.id
                ORDER BY al.timestamp DESC
                LIMIT 5
            """)
        else:
            cursor.execute("""
                SELECT al.activity, al.timestamp, u.full_name
                FROM activity_logs al
                JOIN users u ON al.user_id = u.id
                JOIN user_profiles up ON u.id = up.user_id
                WHERE up.department = %s
                ORDER BY al.timestamp DESC
                LIMIT 5
            """, (user_department,))
            
        recent_activities = cursor.fetchall()

        if is_admin:
            cursor.execute("""
                SELECT DISTINCT department 
                FROM documents 
                WHERE department IS NOT NULL AND department != ''
                ORDER BY department
            """)
        else:
            cursor.execute("""
                SELECT DISTINCT department 
                FROM documents 
                WHERE department = %s AND department IS NOT NULL AND department != ''
                ORDER BY department
            """, (user_department,))
            
        departments_list = cursor.fetchall()
        
        if is_admin:
            cursor.execute("""
                SELECT 
                    d.department AS name,
                    dm.category,
                    COUNT(*) as count
                FROM documents d
                JOIN document_metadata dm ON d.id = dm.document_id
                WHERE d.department IS NOT NULL AND d.department != ''
                GROUP BY d.department, dm.category
                ORDER BY d.department
            """)
        else:
            cursor.execute("""
                SELECT 
                    d.department AS name,
                    dm.category,
                    COUNT(*) as count
                FROM documents d
                JOIN document_metadata dm ON d.id = dm.document_id
                WHERE d.department = %s AND d.department IS NOT NULL AND d.department != ''
                GROUP BY d.department, dm.category
                ORDER BY d.department
            """, (user_department,))
            
        dept_rows = cursor.fetchall()
        
        departments = {}
        
        for dept in departments_list:
            dept_name = dept['department']
            departments[dept_name] = {
                'name': dept_name,
                'counts': {
                    'Copyright': 0,
                    'Patent': 0,
                    'Trademark': 0,
                    'Utility Model': 0,
                    'Industrial Design': 0
                },
                'total': 0
            }
        
        for row in dept_rows:
            dept_name = row['name']
            
            if dept_name not in departments:
                departments[dept_name] = {
                    'name': dept_name,
                    'counts': {
                        'Copyright': 0,
                        'Patent': 0,
                        'Trademark': 0,
                        'Utility Model': 0,
                        'Industrial Design': 0
                    },
                    'total': 0
                }
            
            departments[dept_name]['counts'][row['category']] = row['count']
            departments[dept_name]['total'] += row['count']
        
        department_documents = list(departments.values())

        search_query = request.args.get('search', '').strip()
        search_results = []

        if search_query:
            if is_admin:
                cursor.execute("""
                    SELECT d.title, dm.category, dm.status, dm.submission_date
                    FROM documents d
                    JOIN document_metadata dm ON d.id = dm.document_id
                    WHERE d.title LIKE %s
                    ORDER BY d.created_at DESC
                    LIMIT 10
                """, (f"%{search_query}%",))
            else:
                cursor.execute("""
                    SELECT d.title, dm.category, dm.status, dm.submission_date
                    FROM documents d
                    JOIN document_metadata dm ON d.id = dm.document_id
                    WHERE d.title LIKE %s AND d.department = %s
                    ORDER BY d.created_at DESC
                    LIMIT 10
                """, (f"%{search_query}%", user_department))
                
            search_results = cursor.fetchall()

        return render_template(
            'index.html',
            document_counts=document_counts,
            full_name=full_name,
            recent_activities=recent_activities,
            search_query=search_query,
            search_results=search_results,
            profile=profile,
            department_documents=department_documents,
            is_admin=is_admin,
            user_department=user_department
        )
    
    except Error as e:
        flash('Failed to load dashboard data', 'error')
        print(f"Dashboard error: {e}")
        return redirect(url_for('login'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

@app.route('/chart-data')
def chart_data():
    if 'user_id' not in session:
        return jsonify({"error": "Not authenticated"}), 401
    
    db = get_db_connection()
    if not db:
        return jsonify({"error": "Database connection failed"}), 500
    
    cursor = None
    try:
        cursor = db.cursor(dictionary=True)
        
        user_id = session['user_id']
        
        cursor.execute("""
            SELECT u.*, up.department as user_department 
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return jsonify({"error": "User not found"}), 404
            
        is_admin = user['role'] == 'Admin'
        user_department = user.get('department') or user.get('user_department')
        
        if is_admin:
            cursor.execute("""
                SELECT MONTH(date_encoded) AS month, COUNT(*) AS count
                FROM documents
                GROUP BY MONTH(date_encoded)
                ORDER BY month ASC
            """)
        else:
            cursor.execute("""
                SELECT MONTH(date_encoded) AS month, COUNT(*) AS count
                FROM documents
                WHERE department = %s
                GROUP BY MONTH(date_encoded)
                ORDER BY month ASC
            """, (user_department,))

        data = cursor.fetchall()

        all_months = {f"{i:02d}": 0 for i in range(1, 13)}

        for row in data:
            month_str = f"{row['month']:02d}"
            all_months[month_str] = row['count']

        chart_data = {
            "labels": list(all_months.keys()),
            "values": list(all_months.values())
        }   

        return jsonify(chart_data)
    
    except Error as e:
        print(f"Chart Data Error: {e}")
        return jsonify({"error": "Database Error"}), 500
    finally: 
        if cursor:
            cursor.close()
        if db.is_connected():
            db.close()

@app.route('/statistics-data')
def statistics_data():
    if 'user_id' not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    db = get_db_connection()
    if not db:
        return jsonify({"error": "Database connection failed"}), 500
    
    cursor = None
    try:
        cursor = db.cursor(dictionary=True)
        
        user_id = session['user_id']
        
        cursor.execute("""
            SELECT u.*, up.department as user_department 
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return jsonify({"error": "User not found"}), 404
            
        is_admin = user['role'] == 'Admin'
        user_department = user.get('department') or user.get('user_department')
        
        if is_admin:

            cursor.execute("""
                SELECT category, COUNT(*) as count
                FROM document_metadata
                GROUP BY category
                ORDER BY category
            """)
        else:

            cursor.execute("""
                SELECT dm.category, COUNT(*) as count
                FROM document_metadata dm
                JOIN documents d ON dm.document_id = d.id
                WHERE d.department = %s
                GROUP BY dm.category
                ORDER BY dm.category
            """, (user_department,))

        data = cursor.fetchall()

        categories = [row['category'] for row in data]
        counts = [row['count'] for row in data]

        return jsonify({
            "categories": categories,
            "counts": counts
        })
    
    except Error as e:
        print(f"Statistics Data Error: {e}")
        return jsonify({"error": "Database Error"}), 500
    
    finally:
        if cursor:
            cursor.close()
        if db.is_connected():
            db.close()

@app.route('/department-data')
def department_data():
    if 'user_id' not in session:
        return jsonify({"error": "Not authenticated"}), 401
        
    db = get_db_connection()
    if not db:
        return jsonify({"error": "Database connection failed"}), 500
    
    cursor = None
    try:
        cursor = db.cursor(dictionary=True)
        
        user_id = session['user_id']
        
        cursor.execute("""
            SELECT u.*, up.department as user_department 
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()
        
        if not user:
            return jsonify({"error": "User not found"}), 404
            
        is_admin = user['role'] == 'Admin'
        user_department = user.get('department') or user.get('user_department')
        

        if is_admin:

            cursor.execute("""
                SELECT 
                    CASE 
                        WHEN department = 'College of Business, Admin. & Accountancy' 
                        THEN 'College of Business, Administration & Accountancy'
                        ELSE department
                    END AS department_name,
                    dm.category,
                    COUNT(*) as count
                FROM documents d
                JOIN document_metadata dm ON d.id = dm.document_id
                WHERE d.department IS NOT NULL AND d.department != ''
                GROUP BY department_name, dm.category
                ORDER BY department_name
            """)
        else:

            cursor.execute("""
                SELECT 
                    CASE 
                        WHEN department = 'College of Business, Admin. & Accountancy' 
                        THEN 'College of Business, Administration & Accountancy'
                        ELSE department
                    END AS department_name,
                    dm.category,
                    COUNT(*) as count
                FROM documents d
                JOIN document_metadata dm ON d.id = dm.document_id
                WHERE d.department = %s AND d.department IS NOT NULL AND d.department != ''
                GROUP BY department_name, dm.category
                ORDER BY department_name
            """, (user_department,))
        
        rows = cursor.fetchall()
        
        departments = {}
        for row in rows:
            dept_name = row['department_name']
            category = row['category']
            
            if dept_name not in departments:
                departments[dept_name] = {
                    'name': dept_name,
                    'counts': {
                        'Copyright': 0,
                        'Patent': 0,
                        'Trademark': 0,
                        'Utility Model': 0,
                        'Industrial Design': 0
                    },
                    'total': 0
                }
            
            departments[dept_name]['counts'][category] = row['count']
            departments[dept_name]['total'] += row['count']
        
        return jsonify({
            "departments": list(departments.values())
        })
    
    except Error as e:
        print(f"Department Data Error: {e}")
        return jsonify({"error": "Database Error"}), 500
    
    finally:
        if cursor:
            cursor.close()
        if db.is_connected():
            db.close()
            
@app.route('/status-data')
def get_status_data():
    db = get_db_connection()
    if not db:
        return jsonify({"error": "Database connection failed"}), 500
    
    cursor = None
    try:
        cursor = db.cursor(dictionary=True)

        cursor.execute("""
            SELECT
                DATE_FORMAT(d.date_encoded, '%Y-%m') as month, 
                dm.status,
                COUNT(*) as count
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            GROUP BY DATE_FORMAT(d.date_encoded, '%Y-%m'), dm.status
            ORDER BY month ASC
        """)

        data = cursor.fetchall()

        if not data:
            return jsonify({
                "months": ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"],
                "statuses": {
                    "Approved": [0] * 12,
                    "Pending": [0] * 12,
                    "Rejected": [0] * 12
                }
            })

        months_set = set()
        status_data = {}
        predefined_statuses = ["Approved", "Pending", "Rejected"]

        for row in data:
            month = row['month']
            status = row['status']
            count = row['count']

            months_set.add(month)

            if status not in status_data:
                status_data[status] = {}

            status_data[status][month] = count

        months = sorted(list(months_set))

        month_names = []
        for month in months:
            year, month_num = month.split('-')
            month_num = int(month_num)
            month_names.append(['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec'][month_num-1])

        result_statuses = {status: [] for status in predefined_statuses}

        for status in predefined_statuses:
            for month in months:
                result_statuses[status].append(status_data.get(status, {}).get(month, 0))

        for status, monthly_data in status_data.items():
            if status not in predefined_statuses:
                result_statuses[status] = []
                for month in months:
                    result_statuses[status].append(monthly_data.get(month, 0))

        return jsonify({
            "months": month_names,
            "statuses":result_statuses
        })
    
    except Exception as e:
        print(f"Status Data Error: {e}")
        return jsonify({"error": str(e)}), 500
    
    finally:
        if cursor:
            cursor.close()
        if db and hasattr(db, 'is_connected') and db.is_connected():
            db.close()

@app.route('/coordinator-profile')
def coordinator_profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)
    cursor.execute('SELECT * FROM users WHERE id = %s', (session['user_id'],))
    coordinator = cursor.fetchone()
    db.close()

    return render_template('coordinator_profile.html', coordinator=coordinator)

@app.route('/encode', methods=['GET', 'POST'])
def encode():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)

        user_id = session['user_id']
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None

        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        profile['profile_picture_url'] = url_for('static', filename=f'uploads/profile_pictures/{profile["profile_picture"]}') if profile.get('profile_picture') else url_for('static', filename='avatar.png')
    
        if request.method == 'POST':
            print("Form submitted:", request.form)
            print("Files:", request.files)
            
            title = request.form.get('title')
            document_type = request.form.get('document_type', 'Research')  # Default type if not specified
            
            category = request.form.get('category')
            
            if not title:
                flash('Title is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            
            submission_date = request.form.get('submission_date')
            if not submission_date:
                submission_date = datetime.now().strftime('%Y-%m-%d')
            
            try:
                parsed_date = datetime.strptime(submission_date, '%Y-%m-%d')
                formatted_date = parsed_date.strftime('%Y-%m-%d')
            except ValueError:
                formatted_date = datetime.now().strftime('%Y-%m-%d')
            
            expiration_date = request.form.get('expiration_date')
            
            try:
                cursor.execute(
                    "INSERT INTO documents (title, type, author_id, department, created_at) VALUES (%s, %s, %s, %s, %s)",
                    (title, document_type, user_id, department, formatted_date)
                )
                document_id = cursor.lastrowid
                
                authors = request.form.get('authors', '')
                
                file_path = None
                if 'document_file' in request.files:
                    file = request.files['document_file']
                    if file and file.filename:
                        filename = secure_filename(file.filename)
                        unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                        save_path = os.path.join('static', 'uploads', 'documents', unique_filename)
                        file_path = f"uploads/documents/{unique_filename}"
                        file.save(save_path)
                
                cursor.execute(
                    "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, expiration_date, file_path) VALUES (%s, %s, %s, %s, %s, %s, %s)",
                    (document_id, category, authors, 'Pending', formatted_date, expiration_date, file_path)
                )
                
                db.commit()
                flash('Document has been successfully registered', 'success')
                return redirect(url_for('tables'))
            except Exception as e:
                db.rollback()
                flash(f'Error saving document: {str(e)}', 'error')
                print(f"Error saving document: {e}")
                return render_template('encode.html', profile=profile, full_name=full_name)

        return render_template('encode.html', profile=profile, full_name=full_name)  

    except Error as e:
        flash('Failed to process the form submission', 'error')
        print(f"Encode error: {e}")
        return redirect(url_for('encode'))
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/encode/patent', methods=['GET', 'POST'])
def encode_patent():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']
        
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None

        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        if request.method == 'POST':
            patent_title = request.form.get('patent-title')
            
            if not patent_title:
                flash('Patent title is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            
            cursor.execute(
                "INSERT INTO documents (title, type, author_id, department) VALUES (%s, %s, %s, %s)",
                (patent_title, 'Patent', user_id, department)
            )
            document_id = cursor.lastrowid
            
            applicant_name = f"{request.form.get('applicant-lastname', '')}, {request.form.get('applicant-firstname', '')}"
            if request.form.get('applicant-middlename'):
                applicant_name += f" {request.form.get('applicant-middlename')}"
                
            patent_type = []
            if request.form.get('patent-type') == 'direct':
                patent_type.append('Direct')
            if request.form.get('patent-type') == 'pct':
                patent_type.append('PCT')
            if request.form.get('patent-type') == 'divisional':
                patent_type.append('Divisional')
            if request.form.get('patent-type') == 'priority':
                patent_type.append('With Claim of Priority')
            
            patent_type_str = ', '.join(patent_type) if patent_type else 'Direct'
            
            file_path = None
            if 'document_file' in request.files:
                file = request.files['document_file']
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    save_path = os.path.join('static', 'uploads', 'documents', unique_filename)
                    file_path = f"uploads/documents/{unique_filename}"
                    file.save(save_path)
            
            submission_date = datetime.now().strftime('%Y-%m-%d')
            
            cursor.execute(
                "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, file_path) VALUES (%s, %s, %s, %s, %s, %s)",
                (document_id, 'Patent', applicant_name, 'Pending', submission_date, file_path)
            )
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS patent_metadata (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    document_id INT NOT NULL,
                    patent_type VARCHAR(100),
                    applicant_name VARCHAR(255),
                    applicant_address TEXT,
                    applicant_is_inventor BOOLEAN DEFAULT FALSE,
                    inventor_name VARCHAR(255),
                    agent_name VARCHAR(255),
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """)
            
            applicant_is_inventor = 1 if request.form.get('applicant-is-inventor') else 0
            applicant_address = request.form.get('applicant-address', '')
            
            cursor.execute(
                """INSERT INTO patent_metadata 
                   (document_id, patent_type, applicant_name, applicant_address, applicant_is_inventor) 
                   VALUES (%s, %s, %s, %s, %s)""",
                (document_id, patent_type_str, applicant_name, applicant_address, applicant_is_inventor)
            )
            
            db.commit()
            flash('Patent application has been successfully registered', 'success')
            return redirect(url_for('tables'))
            
        return render_template('encode.html', profile=profile, full_name=full_name)
        
    except Exception as e:
        db.rollback()
        flash(f'Error processing patent application: {str(e)}', 'error')
        print(f"Patent encoding error: {e}")
        return render_template('encode.html', profile=profile, full_name=full_name)
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/view_patent/<int:document_id>', methods=['GET'])
def view_patent(document_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        
        cursor.execute("""
            SELECT d.id, d.title, d.type, d.author_id, d.department, 
                   dm.authors, dm.status, dm.submission_date, dm.file_path
            FROM documents d
            LEFT JOIN document_metadata dm ON d.id = dm.document_id
            WHERE d.id = %s
        """, (document_id,))
        document = cursor.fetchone()
        
        if not document:
            flash('Patent document not found', 'error')
            return redirect(url_for('tables'))
        
        cursor.execute("""
            SELECT * FROM patent_metadata 
            WHERE document_id = %s
        """, (document_id,))
        patent_details = cursor.fetchone()
        
        cursor.execute("""
            SELECT * FROM users 
            WHERE id = %s
        """, (document['author_id'],))
        applicant_user = cursor.fetchone()
        
        patent_info = {
            'document': document,
            'patent_details': patent_details,
            'applicant': applicant_user
        }
        
        return render_template('view_patent.html', patent=patent_info)
    
    except Exception as e:
        flash(f'Error retrieving patent details: {str(e)}', 'error')
        print(f"Patent view error: {e}")
        return redirect(url_for('tables'))
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/encode/utility-model', methods=['GET', 'POST'])
def encode_utility_model():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']
        
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None

        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        if request.method == 'POST':
            model_title = request.form.get('patent-title')
            
            if not model_title:
                flash('Utility model title is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            
            cursor.execute(
                "INSERT INTO documents (title, type, author_id, department) VALUES (%s, %s, %s, %s)",
                (model_title, 'Utility Model', user_id, department)
            )
            document_id = cursor.lastrowid
            
            applicant_name = f"{request.form.get('applicant-lastname', '')}, {request.form.get('applicant-firstname', '')}"
            if request.form.get('applicant-middlename'):
                applicant_name += f" {request.form.get('applicant-middlename')}"
            
            file_path = None
            if 'document_file' in request.files:
                file = request.files['document_file']
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    save_path = os.path.join('static', 'uploads', 'documents', unique_filename)
                    file_path = f"uploads/documents/{unique_filename}"
                    file.save(save_path)
            
            submission_date = datetime.now().strftime('%Y-%m-%d')
            
            cursor.execute(
                "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, file_path) VALUES (%s, %s, %s, %s, %s, %s)",
                (document_id, 'Utility Model', applicant_name, 'Pending', submission_date, file_path)
            )
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS utility_model_metadata (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    document_id INT NOT NULL,
                    model_type VARCHAR(100),
                    applicant_name VARCHAR(255),
                    applicant_address TEXT,
                    applicant_is_creator BOOLEAN DEFAULT FALSE,
                    creator_name VARCHAR(255),
                    technical_field VARCHAR(255),
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """)
            
            applicant_is_creator = 1 if request.form.get('applicant-is-creator') else 0
            applicant_address = request.form.get('applicant-address', '')
            technical_field = request.form.get('technical-field', '')
            model_type = request.form.get('model-type', 'Standard')
            
            cursor.execute(
                """INSERT INTO utility_model_metadata 
                   (document_id, model_type, applicant_name, applicant_address, applicant_is_creator, technical_field) 
                   VALUES (%s, %s, %s, %s, %s, %s)""",
                (document_id, model_type, applicant_name, applicant_address, applicant_is_creator, technical_field)
            )
            
            db.commit()
            flash('Utility model application has been successfully registered', 'success')
            return redirect(url_for('tables'))
            
        return render_template('encode.html', profile=profile, full_name=full_name)
        
    except Exception as e:
        db.rollback()
        flash(f'Error processing utility model application: {str(e)}', 'error')
        print(f"Utility model encoding error: {e}")
        return render_template('encode.html', profile=profile, full_name=full_name)
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/encode/industrial-design', methods=['GET', 'POST'])
def encode_industrial_design():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']
        

        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None


        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        if request.method == 'POST':
            design_title = request.form.get('patent-title')
            
            if not design_title:
                flash('Industrial design title is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            
            cursor.execute(
                "INSERT INTO documents (title, type, author_id, department) VALUES (%s, %s, %s, %s)",
                (design_title, 'Industrial Design', user_id, department)
            )
            document_id = cursor.lastrowid
            
            applicant_name = f"{request.form.get('applicant-lastname', '')}, {request.form.get('applicant-firstname', '')}"
            if request.form.get('applicant-middlename'):
                applicant_name += f" {request.form.get('applicant-middlename')}"
            
            file_path = None
            if 'document_file' in request.files:
                file = request.files['document_file']
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    save_path = os.path.join('static', 'uploads', 'documents', unique_filename)
                    file_path = f"uploads/documents/{unique_filename}"
                    file.save(save_path)
            
            submission_date = datetime.now().strftime('%Y-%m-%d')
            
            cursor.execute(
                "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, file_path) VALUES (%s, %s, %s, %s, %s, %s)",
                (document_id, 'Industrial Design', applicant_name, 'Pending', submission_date, file_path)
            )
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS industrial_design_metadata (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    document_id INT NOT NULL,
                    design_type VARCHAR(100),
                    product_category VARCHAR(255),
                    description TEXT,
                    applicant_name VARCHAR(255),
                    applicant_address TEXT,
                    applicant_is_designer BOOLEAN DEFAULT FALSE,
                    designer_name VARCHAR(255),
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """)
            
            applicant_is_designer = 1 if request.form.get('applicant-is-designer') else 0
            applicant_address = request.form.get('applicant-address', '')
            product_category = request.form.get('product-category', '')
            design_type = request.form.get('design-type', '2D')
            description = request.form.get('design-description', '')
            
            cursor.execute(
                """INSERT INTO industrial_design_metadata 
                   (document_id, design_type, product_category, description, applicant_name, applicant_address, applicant_is_designer) 
                   VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                (document_id, design_type, product_category, description, applicant_name, applicant_address, applicant_is_designer)
            )
            
            db.commit()
            flash('Industrial design application has been successfully registered', 'success')
            return redirect(url_for('tables'))
            
        return render_template('encode.html', profile=profile, full_name=full_name)
        
    except Exception as e:
        db.rollback()
        flash(f'Error processing industrial design application: {str(e)}', 'error')
        print(f"Industrial design encoding error: {e}")
        return render_template('encode.html', profile=profile, full_name=full_name)
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/encode/trademark', methods=['GET', 'POST'])
def encode_trademark():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']
        
        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None


        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:

            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        if request.method == 'POST':

            print("\n=== DEBUG START ===")
            print("Form data received:", request.form)
            print("Files received:", request.files)
            print("Files keys:", list(request.files.keys()))
            

            mark_name = request.form.get('title-mark')
            

            if not mark_name:
                flash('Trademark name is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            

            cursor.execute(
                "INSERT INTO documents (title, type, author_id, department) VALUES (%s, %s, %s, %s)",
                (mark_name, 'Trademark', user_id, department)
            )
            document_id = cursor.lastrowid
            

            applicant_name = f"{request.form.get('applicant-lastname', '')}, {request.form.get('applicant-firstname', '')}"
            if request.form.get('applicant-middlename'):
                applicant_name += f" {request.form.get('applicant-middlename')}"
            

            file_path = None

            if 'mark-image' in request.files:
                uploaded_file = request.files['mark-image']
                

                if uploaded_file and uploaded_file.filename != '':
                    try:
                        print("\n=== FILE UPLOAD DEBUG ===")
                        print("Original filename:", uploaded_file.filename)
                        

                        filename = secure_filename(uploaded_file.filename)
                        print("Secure filename:", filename)
                        

                        unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                        print("Unique filename:", unique_filename)
                        

                        upload_folder = os.path.join(current_app.root_path, 'static', 'uploads', 'trademarks')
                        print("Upload folder path:", upload_folder)
                        
                        os.makedirs(upload_folder, exist_ok=True)
                        print("Folder verified/created")
                        
                        save_path = os.path.join(upload_folder, unique_filename)
                        print("Full save path:", save_path)
                        

                        uploaded_file.save(save_path)
                        print("File saved successfully")
                        

                        if os.path.exists(save_path):
                            print("File verification: EXISTS on disk")
                            file_size = os.path.getsize(save_path)
                            print(f"File size: {file_size} bytes")
                            

                            if file_size > 0:
                                file_path = f"uploads/trademarks/{unique_filename}"
                                print("File path to be stored in DB:", file_path)
                            else:
                                print("File saved but has zero size!")
                                file_path = None
                        else:
                            print("File verification: DOES NOT EXIST on disk")
                            file_path = None
                        
                    except Exception as e:
                        print(f"!!! FILE SAVE ERROR: {str(e)}")
                        traceback.print_exc()  # Print full stack trace for debugging
                        file_path = None
                else:
                    print("Uploaded file exists but has no filename or is empty")
            else:
                print("'mark-image' not found in request.files")
                print("Available keys in request.files:", list(request.files.keys()))


            submission_date = datetime.now().strftime('%Y-%m-%d')
            

            cursor.execute(
                "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, file_path) VALUES (%s, %s, %s, %s, %s, %s)",
                (document_id, 'Trademark', applicant_name, 'Pending', submission_date, file_path)
            )
            

            cursor.execute("""
                CREATE TABLE IF NOT EXISTS trademark_metadata (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    document_id INT NOT NULL,
                    mark_type VARCHAR(100),
                    goods_services TEXT,
                    applicant_name VARCHAR(255),
                    applicant_address TEXT,
                    business_type VARCHAR(100),
                    class_numbers VARCHAR(255),
                    mark_description TEXT,
                    mark_image_path VARCHAR(255),
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """)
            

            applicant_address = request.form.get('applicant-address', '')
            goods_services = request.form.get('goods-services', '')
            business_type = request.form.get('business-type', '')
            mark_type = request.form.get('mark-type', 'Word Mark')
            class_numbers = request.form.get('class-numbers', '')
            mark_description = request.form.get('mark-description', '')
            
            cursor.execute(
                """INSERT INTO trademark_metadata 
                   (document_id, mark_type, goods_services, applicant_name, 
                    applicant_address, business_type, class_numbers, 
                    mark_description, mark_image_path) 
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""",
                (document_id, mark_type, goods_services, applicant_name, 
                 applicant_address, business_type, class_numbers, 
                 mark_description, file_path)
            )


            print("\n=== DATABASE DEBUG ===")
            print("About to insert into trademark_metadata with these values:")
            print(f"document_id: {document_id}")
            print(f"mark_type: {mark_type}")
            print(f"goods_services: {goods_services}")
            print(f"applicant_name: {applicant_name}")
            print(f"file_path: {file_path}")
            print("Insert operation completed")
            print("=== END DATABASE DEBUG ===\n")
            
            db.commit()
            flash('Trademark application has been successfully registered', 'success')
            return redirect(url_for('tables'))
            

        return render_template('encode.html', profile=profile, full_name=full_name)
        
    except Exception as e:
        db.rollback()
        flash(f'Error processing trademark application: {str(e)}', 'error')
        print(f"Trademark encoding error: {e}")
        traceback.print_exc()  # Add this for complete error tracing
        return render_template('encode.html', profile=profile, full_name=full_name)
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/encode/copyright', methods=['GET', 'POST'])
def encode_copyright():
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']
        

        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        department = user['department'] if user else None


        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:

            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        if request.method == 'POST':

            work_title = request.form.get('work-title')
            

            if not work_title:
                flash('Work title is required', 'error')
                return render_template('encode.html', profile=profile, full_name=full_name)
            

            cursor.execute(
                "INSERT INTO documents (title, type, author_id, department) VALUES (%s, %s, %s, %s)",
                (work_title, 'Copyright', user_id, department)
            )
            document_id = cursor.lastrowid
            

            author_name = f"{request.form.get('author-lastname', '')}, {request.form.get('author-firstname', '')}"
            if request.form.get('author-middlename'):
                author_name += f" {request.form.get('author-middlename')}"
            

            file_path = None
            if 'document_file' in request.files:
                file = request.files['document_file']
                if file and file.filename:
                    filename = secure_filename(file.filename)
                    unique_filename = f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{filename}"
                    save_path = os.path.join('static', 'uploads', 'documents', unique_filename)
                    file_path = f"uploads/documents/{unique_filename}"
                    file.save(save_path)
            

            submission_date = datetime.now().strftime('%Y-%m-%d')
            

            cursor.execute(
                "INSERT INTO document_metadata (document_id, category, authors, status, submission_date, file_path) VALUES (%s, %s, %s, %s, %s, %s)",
                (document_id, 'Copyright', author_name, 'Pending', submission_date, file_path)
            )
            

            cursor.execute("""
                CREATE TABLE IF NOT EXISTS copyright_metadata (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    document_id INT NOT NULL,
                    work_type VARCHAR(100),
                    creation_date DATE,
                    first_publication_date DATE,
                    author_name VARCHAR(255),
                    author_is_owner BOOLEAN DEFAULT TRUE,
                    copyright_owner VARCHAR(255),
                    work_description TEXT,
                    FOREIGN KEY (document_id) REFERENCES documents(id) ON DELETE CASCADE
                )
            """)
            

            work_type = request.form.get('work-type', '')
            creation_date = request.form.get('creation-date', '')
            first_publication_date = request.form.get('publication-date', '')
            author_is_owner = 1 if request.form.get('author-is-owner') else 0
            copyright_owner = request.form.get('copyright-owner', author_name) if not author_is_owner else author_name
            work_description = request.form.get('work-description', '')
            

            creation_date = datetime.strptime(creation_date, '%Y-%m-%d').date() if creation_date else None
            first_publication_date = datetime.strptime(first_publication_date, '%Y-%m-%d').date() if first_publication_date else None
            
            cursor.execute(
                """INSERT INTO copyright_metadata 
                   (document_id, work_type, creation_date, first_publication_date, author_name, author_is_owner, copyright_owner, work_description) 
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                (document_id, work_type, creation_date, first_publication_date, author_name, author_is_owner, copyright_owner, work_description)
            )
            
            db.commit()
            flash('Copyright application has been successfully registered', 'success')
            return redirect(url_for('tables'))
            

        return render_template('encode.html', profile=profile, full_name=full_name)
        
    except Exception as e:
        db.rollback()
        flash(f'Error processing copyright application: {str(e)}', 'error')
        print(f"Copyright encoding error: {e}")
        return render_template('encode.html', profile=profile, full_name=full_name)
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()


def standardize_department_name(department):
    """
    Standardizes department names to ensure consistent naming across the system.
    Maps abbreviations and variations to their canonical full names.
    """
    if not department:
        return None
    

    department_mapping = {
        'CCS': 'College of Computer Studies',
        'CCIS': 'College of Computer Studies',
        'CS': 'College of Computer Studies',
        'Computer Studies': 'College of Computer Studies',
        
        'CAS': 'College of Arts and Sciences',
        'Arts & Sciences': 'College of Arts and Sciences',
        
        'CBAA': 'College of Business, Administration & Accountancy',
        'Business': 'College of Business, Administration & Accountancy',
        'Business Admin': 'College of Business, Administration & Accountancy',
        
        'CCJE': 'College of Criminal Justice Education',
        'Criminal Justice': 'College of Criminal Justice Education',
        
        'COF': 'College of Fisheries',
        'Fisheries': 'College of Fisheries',
        
        'CFND': 'College of Food Nutrition and Dietetics',
        'Food & Nutrition': 'College of Food Nutrition and Dietetics',
        'Nutrition': 'College of Food Nutrition and Dietetics',
        
        'CHMT': 'College of Hospitality Management & Tourism',
        'Hospitality': 'College of Hospitality Management & Tourism',
        'Tourism': 'College of Hospitality Management & Tourism',
        
        'CTE': 'College of Teacher Education',
        'Education': 'College of Teacher Education',
        'Teacher Ed': 'College of Teacher Education'
    }
    

    normalized_dept = department.strip().upper()
    for abbr, full_name in department_mapping.items():
        if normalized_dept == abbr.upper():
            return full_name
    

    return department

@app.route('/submit_document', methods=['POST'])
def submit_document():
    db = get_db_connection()
    if not db:
        flash("Database connection failed", "error")
        return redirect(url_for("preview"))
    
    cursor = None

    try:
        cursor = db.cursor()

        personal = session.get('personal_details', {})
        document = session.get('document_details', {})
        department = personal.get('department', '')
        
        filenames = session.get('filenames', [])
        file_paths = session.get('file_paths', [])  
        title = document.get('document_title', '')
        doc_type = document.get('register_type', '')

        cursor.execute("""
            INSERT INTO documents (title, type, date_encoded, department)
            VALUES (%s, %s, NOW(), %s)
        """, (title, doc_type, department))

        print("Personal details:", personal)
        print("Document details:", document)
        print("Filenames:", filenames)

        print("Checking session data completeness...")
        required_fields = {
            'personal': ['firstname', 'lastname'],
            'document': ['document_title', 'registration_type', 'category']
        }

        for category, fields_list in required_fields.items():
            data = personal if category == 'personal' else document
            for field_name in fields_list:
                if field_name not in data or not data[field_name]:
                    print(f"Missing required field: {category}.{field_name}")

        document_id = cursor.lastrowid 

        if not document_id:
            print("Error: Failed to get document ID after insertion")
            flash("Error submitting document: Could not create document record", "error")
            db.rollback()
            return redirect(url_for('preview'))

        category = document.get('category', '')
        status = 'Pending'
        authors = f"{personal.get('firstname', '')} {personal.get('middleinitial', '')} {personal.get('lastname', '')}"
        department = personal.get('department', '')

        submission_date = document.get('submission_date', None)
        expiration_date = document.get('expiration_date', None)

        all_file_paths = ','.join(file_paths) if file_paths else None

        print(f"Inserting metadata with document_id={document_id}, submission_date={submission_date}, expiration_date={expiration_date}")
    
        cursor.execute("""
            INSERT INTO document_metadata (document_id, category, authors, status, submission_date, expiration_date, file_path)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        """, (document_id, category, authors, status, submission_date, expiration_date, all_file_paths))

        log_activity(session['user_id'], f"Created new document: {title} (ID: {document_id})")

        db.commit()
        print("Database commit successful!")
        flash("Document submitted successfully!", "success")

        session.pop('personal_details', None)
        session.pop('document_details', None)
        session.pop('filenames', None)
        session.pop('file_paths', None)
        session.pop('files', None)

        return redirect(url_for('tables'))
    
    except Error as e:
        if db:
            db.rollback()
        print("Error submitting document!", e)
        flash("Failed to submit document.", "error")
        return redirect(url_for('preview'))
    
    finally:
        if cursor:
            cursor.close()
        if db and db.is_connected():
            db.close()

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in {'pdf', 'doc', 'docx', 'txt'}

@app.route('/serve-document/<path:filename>')
def serve_document(filename):
    if ',' in filename:
        filename = filename.split(',')[0]
    
    if filename.startswith('documents\\') or filename.startswith('documents/'):
        filename = filename[10:] 
    
    return send_from_directory(os.path.join(app.static_folder, 'documents'), filename)
    


@app.route('/tables')
def tables():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return render_template('tables.html', profile=None)
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']


        cursor.execute("""
            SELECT u.*, up.department as user_department 
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))        
        user = cursor.fetchone()

        if not user:
            return redirect(url_for('login'))
        
        full_name = user['full_name'] if user else 'User'


        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        profile['profile_picture_url'] = url_for('static', filename=f'uploads/profile_pictures/{profile["profile_picture"]}') if profile.get('profile_picture') else url_for('static', filename='avatar.png')


        status_filter = request.args.get('status', 'all')
        category_filter = request.args.get('category', 'all')
        department_filter = request.args.get('department', 'all')
        

        base_query = """
            SELECT 
                d.id, d.title, dm.category, dm.authors, dm.status, 
                dm.submission_date, dm.expiration_date, d.type, d.department
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            WHERE 1=1
        """
        params = []
        

        user_department = user.get('department') or profile.get('user_department')
        print(f"DEBUG - User Role: {user['role']}, Department: {user_department}")
        
        if user['role'] == 'Coordinator' and user_department:
            base_query += " AND d.department = %s"
            params.append(user_department)
            print(f"DEBUG - Applying department filter: {user_department}")
        else:

            if user['role'] == 'Coordinator':
                print("DEBUG - Coordinator has no department assigned!")
        

        if status_filter != 'all':
            base_query += " AND dm.status = %s"
            params.append(status_filter)
            
        if category_filter != 'all':
            base_query += " AND dm.category = %s"
            params.append(category_filter)
            
        if department_filter != 'all':
            base_query += " AND d.department = %s"
            params.append(department_filter)
        
        base_query += " ORDER BY d.id DESC"
        

        print(f"DEBUG - Final Query: {base_query}")
        print(f"DEBUG - Parameters: {params}")

        cursor.execute(base_query, params)
        documents = cursor.fetchall()
        print(f"DEBUG - Query returned {len(documents)} documents")
        

        if not documents:
            cursor.execute("SELECT COUNT(*) as count FROM documents")
            total_docs = cursor.fetchone()['count']
            print(f"DEBUG - Total documents in database: {total_docs}")
            
            if total_docs > 0:

                cursor.execute("SELECT COUNT(*) as count FROM documents d JOIN document_metadata dm ON d.id = dm.document_id")
                joined_docs = cursor.fetchone()['count']
                print(f"DEBUG - Total joined documents: {joined_docs}")
                

                if joined_docs > 0:
                    cursor.execute("SELECT d.id, d.title, d.department FROM documents d LIMIT 5")
                    sample_docs = cursor.fetchall()
                    print("DEBUG - Sample documents:")
                    for doc in sample_docs:
                        print(f"  ID: {doc['id']}, Title: {doc['title']}, Department: {doc['department']}")

        cursor.execute("SELECT DISTINCT category FROM document_metadata")
        categories = cursor.fetchall()


        cursor.execute("SELECT DISTINCT department FROM documents WHERE department IS NOT NULL AND department != ''")
        department_results = cursor.fetchall()
        departments = [dept['department'] for dept in department_results]
        

        print(f"DEBUG - Found departments: {departments}")

        return render_template('tables.html', 
                               documents=documents,
                               categories=categories,
                               departments=departments,  # Pass all departments to template
                               selected_status=status_filter,
                               selected_category=category_filter,
                               selected_department=department_filter,
                               profile=profile,
                               user_role=user['role'],
                               user_department=user_department,
                               full_name=full_name,
                               is_admin=(user['role'] == 'Admin'))

    except Error as e:
        flash('Failed to fetch documents data', 'error')
        print(f"Error fetching documents: {e}")
        return render_template('tables.html', documents=[], profile=None)

    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/document-status-update', methods=['POST'])
def update_status():
    if request.method == 'POST':
        data = request.json
        document_id = data.get('document_id')
        new_status = data.get('status')
        
        if not document_id or not new_status:
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        
        try:
            db = get_db_connection()
            if not db:
                return jsonify({'success': False, 'message': 'Database connection failed'}), 500
            
            cursor = db.cursor()
            cursor.execute("""
                UPDATE document_metadata 
                SET status = %s 
                WHERE document_id = %s
            """, (new_status, document_id))
            
            db.commit()
            
            if cursor.rowcount > 0:
                return jsonify({'success': True, 'message': 'Status updated successfully'})
            else:
                return jsonify({'success': False, 'message': 'Document not found or status unchanged'}), 404
                
        except Error as e:
            print(f"Database error: {e}")
            return jsonify({'success': False, 'message': f'Database error: {str(e)}'}), 500
            
        finally:
            if db.is_connected():
                cursor.close()
                db.close()
    
    return jsonify({'success': False, 'message': 'Invalid request method'}), 405


@app.route('/report')
def report():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return render_template('tables.html', profile=None)

    user_id = session['user_id']
    cursor = db.cursor(dictionary=True)

    try:

        cursor.execute("""
            SELECT u.*, up.first_name, up.last_name, up.email as profile_email, 
                   up.phone, up.position, up.institution, up.office_address, 
                   up.profile_picture, up.department as profile_department
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (user_id,))
        user = cursor.fetchone()

        if not user:
            flash('User not found', 'error')
            return redirect(url_for('login'))


        profile = {
            'user_id': user_id,
            'first_name': user.get('first_name') or user.get('full_name', '').split(' ')[0],
            'last_name': user.get('last_name') or ' '.join(user.get('full_name', '').split(' ')[1:]),
            'email': user.get('profile_email') or user.get('email'),
            'phone': user.get('phone', ''),
            'position': user.get('position', 'Coordinator'),
            'institution': user.get('institution', ''),
            'office_address': user.get('office_address', ''),
            'profile_picture': user.get('profile_picture', 'avatar.png'),
            'role': user.get('role', 'Coordinator').lower(),
            'department': user.get('profile_department') or user.get('department')
        }


        user_role = 'admin' if user.get('role', '').lower() == 'admin' or profile['role'] == 'admin' else 'coordinator'
        user_department = profile['department']


        base_query = """
            SELECT 
                dm.category,
                COUNT(*) as total_documents,
                SUM(CASE WHEN dm.expiration_date < NOW() THEN 1 ELSE 0 END) as expired_documents,
                SUM(CASE WHEN dm.status = 'Approved' THEN 1 ELSE 0 END) as approved,
                SUM(CASE WHEN dm.status = 'Pending' THEN 1 ELSE 0 END) as pending,
                SUM(CASE WHEN dm.status = 'Rejected' THEN 1 ELSE 0 END) as rejected
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            {department_filter}
            GROUP BY dm.category
        """


        department_filter = ""
        params = ()

        if user_role == 'coordinator' and user_department:
            department_filter = "WHERE d.department = %s"  # Change dm to d here
            params = (user_department,)

        cursor.execute(base_query.format(department_filter=department_filter), params)
        categories_data = cursor.fetchall()


        categories = []
        for cat in categories_data:
            total = cat['total_documents']
            approved = cat['approved']
            rejected = cat['rejected']
            
            categories.append({
                'name': cat['category'],
                'total_documents': total,
                'expired_documents': cat['expired_documents'],
                'approved': approved,
                'pending': cat['pending'],
                'rejected': rejected,
                'approval_rate': round((approved / total) * 100, 1) if total > 0 else 0,
                'rejection_rate': round((rejected / total) * 100, 1) if total > 0 else 0,
                'avg_processing_time': 7.5,  # Replace with actual calculation
            })

        chart_data = {
            'approved': sum(cat['approved'] for cat in categories),
            'pending': sum(cat['pending'] for cat in categories),
            'rejected': sum(cat['rejected'] for cat in categories),
            'categories': [cat['name'] for cat in categories],
            'approvalRate': [cat['approval_rate'] for cat in categories],
            'rejectionRate': [cat['rejection_rate'] for cat in categories],
            'approvedByCategory': [cat['approved'] for cat in categories],
            'pendingByCategory': [cat['pending'] for cat in categories],
            'rejectedByCategory': [cat['rejected'] for cat in categories],
            'user_role': user_role
        }

        return render_template('report.html', 
                            categories=categories, 
                            chart_data=chart_data, 
                            profile=profile, 
                            full_name=f"{profile['first_name']} {profile['last_name']}",
                            user_role=user_role)

    except Exception as e:
        flash('Failed to load report data', 'error')
        print(f"Error loading report data: {e}")
        return render_template('report.html', 
                            profile=None, 
                            full_name='')

    finally:
        if cursor:
            cursor.close()
        if db and db.is_connected():
            db.close()

@app.route('/api/filter-data', methods=['POST'])
def filter_data():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    data = request.get_json()
    db = get_db_connection()
    cursor = db.cursor(dictionary=True)

    try:

        cursor.execute("""
            SELECT u.role, COALESCE(up.department, u.department) as department
            FROM users u
            LEFT JOIN user_profiles up ON u.id = up.user_id
            WHERE u.id = %s
        """, (session['user_id'],))
        user_info = cursor.fetchone()

        if not user_info:
            return jsonify({'error': 'User not found'}), 404

        user_role = user_info['role'].lower()
        user_department = user_info['department']


        base_query = """
            SELECT 
                dm.category,
                COUNT(*) as total_documents,
                SUM(CASE WHEN dm.expiration_date < NOW() THEN 1 ELSE 0 END) as expired_documents,
                SUM(CASE WHEN dm.status = 'Approved' THEN 1 ELSE 0 END) as approved,
                SUM(CASE WHEN dm.status = 'Pending' THEN 1 ELSE 0 END) as pending,
                SUM(CASE WHEN dm.status = 'Rejected' THEN 1 ELSE 0 END) as rejected
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            WHERE 1=1
        """

        params = []
        

        if user_role == 'coordinator' and user_department:
            base_query += " AND dm.department = %s"
            params.append(user_department)
        elif data.get('department') and data['department'] != 'All' and user_role == 'admin':
            base_query += " AND dm.department = %s"
            params.append(data['department'])


        if data.get('startDate'):
            base_query += " AND dm.created_at >= %s"
            params.append(data['startDate'])
        if data.get('endDate'):
            base_query += " AND dm.created_at <= %s"
            params.append(data['endDate'])
        if data.get('status') and data['status'] != 'All':
            base_query += " AND dm.status = %s"
            params.append(data['status'])
        if data.get('category') and data['category'] != 'All':
            base_query += " AND dm.category = %s"
            params.append(data['category'])

        base_query += " GROUP BY dm.category"

        cursor.execute(base_query, params)
        categories_data = cursor.fetchall()


        categories = []
        for cat in categories_data:
            total = cat['total_documents']
            approved = cat['approved']
            rejected = cat['rejected']
            
            categories.append({
                'name': cat['category'],
                'total_documents': total,
                'expired_documents': cat['expired_documents'],
                'approved': approved,
                'pending': cat['pending'],
                'rejected': rejected,
                'approval_rate': round((approved / total) * 100, 1) if total > 0 else 0,
                'rejection_rate': round((rejected / total) * 100, 1) if total > 0 else 0,
                'avg_processing_time': 7.5,
            })

        chart_data = {
            'tableData': categories,
            'chartData': {
                'approved': sum(cat['approved'] for cat in categories),
                'pending': sum(cat['pending'] for cat in categories),
                'rejected': sum(cat['rejected'] for cat in categories),
                'categories': [cat['name'] for cat in categories],
                'approvalRate': [cat['approval_rate'] for cat in categories],
                'rejectionRate': [cat['rejection_rate'] for cat in categories],
                'approvedByCategory': [cat['approved'] for cat in categories],
                'pendingByCategory': [cat['pending'] for cat in categories],
                'rejectedByCategory': [cat['rejected'] for cat in categories]
            }
        }

        return jsonify(chart_data)

    except Exception as e:
        print(f"Error filtering data: {e}")
        return jsonify({'error': 'Failed to filter data'}), 500

    finally:
        if cursor:
            cursor.close()
        if db and db.is_connected():
            db.close()

@app.route('/api/generate-report', methods=['POST'])
def generate_report_data():
    if 'user_id' not in session:
        return jsonify({"error": "Not logged in"}), 401
    
    try:
        response = filter_data()
        return response
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        return jsonify({"error": "Report generation failed"}), 500
    
@app.route("/generate_report")
def generate_report():
    response = make_response()
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = 'attachment; filename=report.pdf'

    pdf = canvas.Canvas(response, pagesize=letter)
    pdf.drawString(100, 750, "Your PDF Report")
    pdf.drawString(100, 730, "Data: Example report data goes here")

    pdf.showPage()
    pdf.save()

    return response

@app.route('/generate_word_report', methods=['GET', 'POST'])
def generate_word_report():
    if 'user_id' not in session:
        return jsonify({"error": "Not logged in"}), 401
    

    print("Starting Word report generation process")
    
    try:

        try:
            from docx import Document
            from docx.shared import Inches
            from io import BytesIO
            import base64
            print("Successfully imported libraries")
        except ImportError as e:
            print(f"Error importing libraries: {str(e)}")
            return jsonify({"error": "Missing required libraries. Please install python-docx."}), 500
        

        chart_data = None
        table_data = None
        
        if request.method == 'POST':
            try:
                data = request.get_json()
                if data:
                    print("Received POST data with chart images")
                    line_chart_b64 = data.get('lineChartImage')
                    pie_chart_b64 = data.get('pieChartImage')
                    bar_chart_b64 = data.get('barChartImage')
                    table_data = data.get('tableData')
                    
                    chart_data = {
                        'line': line_chart_b64,
                        'pie': pie_chart_b64,
                        'bar': bar_chart_b64
                    }
            except Exception as json_error:
                print(f"Error processing JSON data: {str(json_error)}")
                return jsonify({"error": "Invalid JSON data"}), 400
        

        if not table_data:
            try:
                db = get_db_connection()
                if not db:
                    print("Database connection failed")
                    return jsonify({"error": "Database connection failed"}), 500
                print("Successfully connected to database")
                

                cursor = db.cursor(dictionary=True)
                print("Database cursor created")
                

                user_id = session.get('user_id')
                user_role = session.get('role', 'coordinator')
                user_department = session.get('department')
                print(f"User info - ID: {user_id}, Role: {user_role}, Department: {user_department}")
                

                base_query = """
                    SELECT 
                        dm.category,
                        COUNT(*) as total_documents,
                        SUM(CASE WHEN dm.expiration_date < NOW() THEN 1 ELSE 0 END) as expired_documents,
                        SUM(CASE WHEN dm.status = 'Approved' THEN 1 ELSE 0 END) as approved,
                        SUM(CASE WHEN dm.status = 'Pending' THEN 1 ELSE 0 END) as pending,
                        SUM(CASE WHEN dm.status = 'Rejected' THEN 1 ELSE 0 END) as rejected
                    FROM documents d
                    JOIN document_metadata dm ON d.id = dm.document_id
                    {department_filter}
                    GROUP BY dm.category
                """
                
                department_filter = ""
                params = ()
                
                if user_role == 'coordinator' and user_department:
                    department_filter = "WHERE d.department = %s"
                    params = (user_department,)
                    

                print(f"Executing query with params: {params}")
                cursor.execute(base_query.format(department_filter=department_filter), params)
                table_data = cursor.fetchall()
                print(f"Query successful, retrieved {len(table_data)} category records")
                

                if not table_data:
                    print("No data retrieved from database")
                    return jsonify({"error": "No report data available"}), 404
                    

                for cat in table_data:
                    total = cat['total_documents']
                    approved = cat['approved']
                    rejected = cat['rejected']
                    
                    cat['approval_rate'] = round((approved / total) * 100, 1) if total > 0 else 0
                    cat['rejection_rate'] = round((rejected / total) * 100, 1) if total > 0 else 0
                    cat['avg_processing_time_a'] = "7.5 days"  # Replace with actual calculation
                    cat['avg_processing_time_r'] = "7.5 days"  # Replace with actual calculation
                
            except Exception as db_error:
                print(f"Database error: {str(db_error)}")
                return jsonify({"error": f"Database error: {str(db_error)}"}), 500
                
            finally:
                if 'cursor' in locals() and cursor:
                    print("Closing cursor")
                    cursor.close()
                if 'db' in locals() and db and hasattr(db, 'is_connected') and db.is_connected():
                    print("Closing database connection")
                    db.close()
        

        try:
            print("Creating Word document")
            doc = Document()
            

            core_properties = doc.core_properties
            core_properties.title = 'Intellectual Property Report'
            core_properties.category = 'Reports'
            core_properties.content_status = 'Final'
            
            doc.add_heading('Intellectual Property Report', 0)
            
            from datetime import datetime
            doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            

            print("Adding table to document")
            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            headers = ['Category', 'Total Submitted', 'Total Expired', 'Approved', 
                      'Pending', 'Rejected', 'Approval Rate', 'Rejection Rate', 
                      'Avg. Processing Time (A)', 'Avg. Processing Time (R)']
            
            for i, header in enumerate(headers):
                header_cells[i].text = header
            

            print("Populating table with data")
            for cat in table_data:
                row_cells = table.add_row().cells
                

                if isinstance(cat, dict):
                    row_cells[0].text = str(cat.get('category', ''))
                    row_cells[1].text = str(cat.get('total_documents', 0))
                    row_cells[2].text = str(cat.get('expired_documents', 0))
                    row_cells[3].text = str(cat.get('approved', 0))
                    row_cells[4].text = str(cat.get('pending', 0))
                    row_cells[5].text = str(cat.get('rejected', 0))
                    row_cells[6].text = f"{cat.get('approval_rate', 0)}%"
                    row_cells[7].text = f"{cat.get('rejection_rate', 0)}%"
                    row_cells[8].text = str(cat.get('avg_processing_time_a', "N/A"))
                    row_cells[9].text = str(cat.get('avg_processing_time_r', "N/A"))
            

            if chart_data:
                doc.add_heading('Charts', level=1)
                

                if chart_data.get('line'):
                    doc.add_heading('Average Approval and Rejection Rate', level=2)
                    line_image_data = base64.b64decode(chart_data['line'])
                    line_image_stream = BytesIO(line_image_data)
                    doc.add_picture(line_image_stream, width=Inches(6))
                

                if chart_data.get('pie'):
                    doc.add_heading('Document Status Distribution', level=2)
                    pie_image_data = base64.b64decode(chart_data['pie'])
                    pie_image_stream = BytesIO(pie_image_data)
                    doc.add_picture(pie_image_stream, width=Inches(5))
                

                if chart_data.get('bar'):
                    doc.add_heading('Comparison of Status by Category', level=2)
                    bar_image_data = base64.b64decode(chart_data['bar'])
                    bar_image_stream = BytesIO(bar_image_data)
                    doc.add_picture(bar_image_stream, width=Inches(6))
            else:
                doc.add_heading('Charts', level=1)
                doc.add_paragraph('For detailed charts, please refer to the online dashboard.')
            

            doc.add_paragraph().add_run('\nThis document is fully editable.').bold = True
            

            print("Saving document to memory")
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            print("Successfully generated Word report")
            

            response = send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'ip_report_{datetime.now().strftime("%Y%m%d")}.docx'
            )
            

            response.headers['X-Content-Type-Options'] = 'nosniff'
            response.headers['Content-Security-Policy'] = "default-src 'self'"
            
            return response
            
        except Exception as doc_error:
            print(f"Document creation error: {str(doc_error)}")
            return jsonify({"error": f"Word document creation failed: {str(doc_error)}"}), 500
    
    except Exception as e:
        import traceback
        print(f"Error generating Word report: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": "Word report generation failed", "details": str(e)}), 500

@app.route('/verify_user', methods=['POST'])
def verify_user():
    print("Route /verify_user accessed")
    conn = None  
    try:
        user_id = request.json.get('user_id')
        department = request.json.get('department')  # Optional department update
        
        if not user_id:
            return jsonify({'success': False, 'message': 'User ID missing'}), 400


        current_user_id = session.get('user_id')
        if not current_user_id:
            return jsonify({'success': False, 'message': 'You must be logged in to verify users'}), 401

        conn = get_db_connection()
        cursor = conn.cursor()


        if department:

            cursor.execute("UPDATE users SET is_verified = 1, department = %s WHERE id = %s", 
                         (department, user_id))
        else:

            cursor.execute("UPDATE users SET is_verified = 1 WHERE id = %s", (user_id,))
        

        cursor.execute("SELECT * FROM user_management WHERE user_id = %s", (user_id,))
        existing_record = cursor.fetchone()
        
        if existing_record:

            cursor.execute("""
                UPDATE user_management 
                SET status = 'Verified' 
                WHERE user_id = %s
            """, (user_id,))
        else:

            cursor.execute("""
                INSERT INTO user_management (user_id, status, created_by) 
                VALUES (%s, 'Verified', %s)
            """, (user_id, current_user_id))
        
        if cursor.rowcount == 0:
            return jsonify({'success': False, 'message': 'User not found'}), 404
            
        conn.commit()
        return jsonify({'success': True, 'message': 'User verified successfully!'})

    except Exception as e:
        print(f"Error: {str(e)}")
        if conn:
            conn.rollback()
        return jsonify({'success': False, 'message': str(e)}), 500

    finally:
        if conn:
            conn.close()

@app.route('/user', methods=['GET'])
def user():
    if 'user_id' not in session:
        return redirect(url_for('login'))
        
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']


        cursor.execute('SELECT * FROM users WHERE id = %s', (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        user_role = user['role'] if user else 'Coordinator'  # Get user role for permission check
        department = user['department'] if user else None

        cursor.execute('SELECT * FROM user_profiles WHERE user_id = %s', (user_id,))
        profile = cursor.fetchone()

        if not profile:

            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': user.get('role', 'Coordinator'),
                'department': user.get('department', ''),  # Use the department from users table
                'institution': user.get('department', ''),  # Keep institution for backward compatibility
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()
        else:

            profile['department'] = user.get('department', '')
            

        all_users = []
        if user_role == 'Admin':

            cursor.execute("""
                SELECT 
                    u.id, 
                    u.full_name, 
                    u.email, 
                    u.role as position, 
                    u.department,
                    u.is_verified,
                    COALESCE(um.status, CASE WHEN u.is_verified THEN 'active' ELSE 'pending' END) as user_status
                FROM users u
                LEFT JOIN user_management um ON u.id = um.user_id
                WHERE u.role = 'Coordinator'
                ORDER BY u.full_name
            """)
            all_users = cursor.fetchall()
            

            for u in all_users:
                if u['user_status'] == 'suspended':
                    u['status'] = 'Paused'
                elif u['is_verified']:
                    u['status'] = 'Verified'
                else:
                    u['status'] = 'Pending'
                

                if u['department'] is None:
                    u['department'] = 'None'

        if profile.get('profile_picture'):
            profile['profile_picture_url'] = url_for('static', filename=f'uploads/profile_pictures/{profile["profile_picture"]}')
        else:
            profile['profile_picture_url'] = url_for('static', filename='avatar.png')


        return render_template('user.html', 
                              profile=profile, 
                              full_name=full_name, 
                              all_users=all_users, 
                              department=department, 
                              user_role=user_role)
    
    except Error as e:
        flash(f'Error loading profile: {str(e)}', 'error')
        return redirect(url_for('dashboard'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

@app.route('/pause_user', methods=['POST'])
def pause_user():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401
    
    data = request.get_json()
    if not data or 'user_id' not in data:
        return jsonify({'success': False, 'message': 'Missing user_id parameter'}), 400
    
    user_id = data['user_id']
    admin_id = session['user_id']  # The admin performing the action
    
    db = get_db_connection()
    if not db:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500

    cursor = db.cursor()
    
    try:

        cursor.execute("SELECT id FROM users WHERE id = %s", (user_id,))
        if not cursor.fetchone():
            return jsonify({'success': False, 'message': 'User not found'}), 404
        

        cursor.execute("SELECT id FROM user_management WHERE user_id = %s", (user_id,))
        if cursor.fetchone():

            cursor.execute("""
                UPDATE user_management 
                SET status = 'suspended', 
                    updated_at = CURRENT_TIMESTAMP 
                WHERE user_id = %s
            """, (user_id,))
        else:

            cursor.execute("""
                INSERT INTO user_management 
                (user_id, status, created_by) 
                VALUES (%s, 'suspended', %s)
            """, (user_id, admin_id))
        
        db.commit()
        

        cursor.execute("SELECT status FROM user_management WHERE user_id = %s", (user_id,))
        result = cursor.fetchone()
        
        if not result or result[0] != 'suspended':
            raise Exception("Status not updated correctly")
            
        return jsonify({'success': True, 'message': 'User suspended successfully'})
    
    except Exception as e:
        db.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
    finally:
        cursor.close()
        db.close()

@app.route('/delete_user', methods=['POST'])
def delete_user():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401
    
    data = request.get_json()
    if not data or 'user_id' not in data:
        return jsonify({'success': False, 'message': 'Missing user_id parameter'}), 400
    
    user_id = data['user_id']
    
    db = get_db_connection()
    if not db:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        cursor = db.cursor()
        

        cursor.execute("DELETE FROM user_profiles WHERE user_id = %s", (user_id,))
        

        cursor.execute("DELETE FROM user_management WHERE user_id = %s", (user_id,))
        

        cursor.execute("DELETE FROM users WHERE id = %s", (user_id,))
        
        db.commit()
        
        return jsonify({'success': True, 'message': 'User deleted successfully'})
    
    except Error as e:
        db.rollback()  # Roll back any changes if an error occurs
        return jsonify({'success': False, 'message': f'Database error: {str(e)}'}), 500
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/reactivate_user', methods=['POST'])
def reactivate_user():
    """
    Route to reactivate a suspended user
    """
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401
    
    data = request.get_json()
    if not data or 'user_id' not in data:
        return jsonify({'success': False, 'message': 'Missing user_id parameter'}), 400
    
    user_id = data['user_id']
    admin_id = session['user_id']  # The admin performing the action is required for created_by
    
    db = get_db_connection()
    if not db:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
        
    cursor = db.cursor()
    
    try:

        cursor.execute("SELECT id FROM users WHERE id = %s", (user_id,))
        if not cursor.fetchone():
            return jsonify({'success': False, 'message': 'User not found'}), 404
        

        cursor.execute("SELECT id, status FROM user_management WHERE user_id = %s", (user_id,))
        management_record = cursor.fetchone()
        
        if management_record:


            cursor.execute("""
                UPDATE user_management 
                SET status = 'active'
                WHERE user_id = %s
            """, (user_id,))
        else:

            cursor.execute("""
                INSERT INTO user_management 
                (user_id, status, created_by) 
                VALUES (%s, 'active', %s)
            """, (user_id, admin_id))
        

        cursor.execute("""
            UPDATE users
            SET is_verified = 1
            WHERE id = %s
        """, (user_id,))
        
        db.commit()
        
        return jsonify({'success': True, 'message': 'User reactivated successfully'})
    
    except Exception as e:
        db.rollback()
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500
    finally:
        cursor.close()
        db.close()

@app.route('/edit_user', methods=['POST'])
def edit_user():
    if 'user_id' not in session:
        return jsonify({'success': False, 'message': 'Not authenticated'}), 401
    
    data = request.get_json()
    if not data or 'user_id' not in data:
        return jsonify({'success': False, 'message': 'Missing user_id parameter'}), 400
    
    user_id = data['user_id']

    full_name = data.get('full_name')
    email = data.get('email')
    department = data.get('department')
    role = data.get('role')
    
    db = get_db_connection()
    if not db:
        return jsonify({'success': False, 'message': 'Database connection failed'}), 500
    
    try:
        cursor = db.cursor()
        

        update_fields = []
        update_values = []
        
        if full_name:
            update_fields.append("full_name = %s")
            update_values.append(full_name)
        
        if email:
            update_fields.append("email = %s")
            update_values.append(email)
        
        if department:
            update_fields.append("department = %s")
            update_values.append(department)
        
        if role:
            update_fields.append("role = %s")
            update_values.append(role)
        
        if not update_fields:
            return jsonify({'success': False, 'message': 'No fields to update'}), 400
        

        update_values.append(user_id)
        

        query = f"UPDATE users SET {', '.join(update_fields)} WHERE id = %s"
        cursor.execute(query, tuple(update_values))
        db.commit()
        
        if cursor.rowcount == 0:
            return jsonify({'success': False, 'message': 'User not found or no changes made'}), 404
        
        return jsonify({'success': True, 'message': 'User updated successfully'})
    
    except Error as e:
        return jsonify({'success': False, 'message': f'Database error: {str(e)}'}), 500
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

@app.route('/update-profile', methods=['POST'])
def update_profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('user'))
    
    try:
        cursor = db.cursor(dictionary=True)
        user_id = session['user_id']

        first_name = request.form.get('first_name', '')
        last_name = request.form.get('last_name', '')
        email = request.form.get('email', '')
        phone = request.form.get('phone', '')
        position = request.form.get('position', '')
        department = request.form.get('department', '')
        office_address = request.form.get('office_address', '')


        cursor.execute("""
            UPDATE user_profiles SET 
                first_name = %s,
                last_name = %s, 
                email = %s,
                phone = %s,
                position = %s,
                institution = %s,  # This is still 'institution' in the database
                office_address = %s
            WHERE user_id = %s
        """, (first_name, last_name, email, phone, position, department, office_address, user_id))
        

        cursor.execute("UPDATE users SET email = %s, department = %s WHERE id = %s", 
                      (email, department, user_id))

        full_name = f"{first_name} {last_name}".strip()
        cursor.execute("UPDATE users SET full_name = %s WHERE id = %s", (full_name, user_id))

        db.commit()
        flash('Profile updated successfully!', 'success')

    except Error as e:
        db.rollback()
        flash(f'Error updating profile: {str(e)}', 'error')

    finally:
        if db.is_connected():
            cursor.close()
            db.close()

    return redirect(url_for('user'))

@app.route('/get_users')
def get_users():
    if 'user_id' not in session:
        return jsonify({'error': 'Unauthorized'}), 401
    
    try:
        db = get_db_connection()
        cursor = db.cursor(dictionary=True)
        cursor.execute("""
            SELECT u.id, u.full_name, u.email, u.auth_code, u.role, 
                   p.position, p.institution, um.status
            FROM users u
            LEFT JOIN user_profiles p ON u.id = p.user_id
            LEFT JOIN user_management um ON u.id = um.user_id
            WHERE u.role = 'coordinator'
            ORDER BY u.full_name
        """)
        users = cursor.fetchall()
        return jsonify(users)
    except Error as e:
        return jsonify({'error': str(e)}), 500
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/upload_profile_picture', methods=['POST'])
def upload_profile_picture():
    if 'user_id' not in session:
        return jsonify({'success': False, 'error': 'Not logged in'})
    
    if 'profile_picture' not in request.files:
        return jsonify({'success': False, 'error': 'No file selected'})
    
    file = request.files['profile_picture']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'})
    
    if file and allowed_file(file.filename):
        try:
            db = get_db_connection()
            if not db:
                return jsonify({'success': False, 'error': 'Database connection failed'})
            
            cursor = db.cursor()
            user_id = session['user_id']
            
            ext = file.filename.rsplit('.', 1)[1].lower()
            filename = secure_filename(f"user_{user_id}.{ext}")
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            file.save(filepath)
            
            cursor.execute("""
                UPDATE user_profiles 
                SET profile_picture = %s 
                WHERE user_id = %s
            """, (filename, user_id))
            db.commit()
            
            log_activity(user_id, "Updated profile picture")
            
            return jsonify({
                'success': True, 
                'filename': filename,
                'url': url_for('static', filename=f'uploads/profile_pictures/{filename}')
            })
            
        except Error as e:
            db.rollback()
            return jsonify({'success': False, 'error': str(e)})
            
        finally:
            if db and db.is_connected():
                cursor.close()
                db.close()
    
    return jsonify({'success': False, 'error': 'Invalid file type'})

def log_activity(user_id, activity, department=None):
    try:
        conn = get_db_connection()
        if not conn:
            print("Failed to log activity: Database Connection Failed")
            return False
        

        if department is None:
            cursor = conn.cursor(dictionary=True)
            cursor.execute("SELECT department FROM users WHERE id = %s", (user_id,))
            user_data = cursor.fetchone()
            if user_data and user_data['department']:
                department = user_data['department']
        
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO activity_logs(user_id, activity, department, timestamp)
            VALUES (%s, %s, %s, %s)
        """, (user_id, activity, department, datetime.now()))
        conn.commit()
        print(f"Activity logged successfully: {activity}")
        return True
    
    except Error as e:
        print(f"Error logging activity: {e}")
        return False

    finally:
        if cursor:
            cursor.close()
        if conn and conn.is_connected():
            conn.close()

def get_recent_activities(user_id, limit=20):
    try:
        conn = get_db_connection()
        if not conn:
            print("Failed to retrieve activities: Database Connection Failed")
            return []
        
        cursor = conn.cursor(dictionary=True)
        

        cursor.execute("SELECT role, department FROM users WHERE id = %s", (user_id,))
        user_data = cursor.fetchone()
        
        if not user_data:
            return []
        

        if user_data['role'] == 'Admin':
            cursor.execute("""
                SELECT al.id, al.activity, al.department, al.timestamp, u.full_name 
                FROM activity_logs al
                JOIN users u ON al.user_id = u.id
                ORDER BY al.timestamp DESC
                LIMIT %s
            """, (limit,))

        elif user_data['role'] == 'Coordinator' and user_data['department']:
            cursor.execute("""
                SELECT al.id, al.activity, al.department, al.timestamp, u.full_name 
                FROM activity_logs al
                JOIN users u ON al.user_id = u.id
                WHERE al.department = %s
                ORDER BY al.timestamp DESC
                LIMIT %s
            """, (user_data['department'], limit))
        else:
            return []
        
        activities = cursor.fetchall()
        return activities
    
    except Error as e:
        print(f"Error retrieving activities: {e}")
        return []

    finally:
        if cursor:
            cursor.close()
        if conn and conn.is_connected():
            conn.close()

def display_recent_activities(activities):
    """Format and display activity logs"""
    if not activities:
        return "<p>No recent activities found</p>"
    
    html = "<div class='activity-logs'>"
    html += "<h3>Recent Activities</h3>"
    html += "<ul class='list-group'>"
    
    for activity in activities:
        formatted_time = activity['timestamp'].strftime("%Y-%m-%d %H:%M:%S")
        html += f"""
        <li class='list-group-item'>
            <div class='d-flex justify-content-between'>
                <span><strong>{activity['full_name']}</strong>: {activity['activity']}</span>
                <span class='text-muted'>{formatted_time}</span>
            </div>
            <small class='text-primary'>Department: {activity['department'] or 'N/A'}</small>
        </li>
        """
    
    html += "</ul></div>"
    return html

@app.route('/update-document/<int:document_id>', methods=['POST'])
def update_document(document_id):
    if 'user_id' not in session:
        return jsonify({"success": False, "error": "Not logged in"}), 401
    
    db = get_db_connection()
    if not db:
        return jsonify({"success": False, "error": "Database connection failed"}), 500
    
    try:

        data = request.json
        

        title = data.get('title')
        category = data.get('category')
        authors = data.get('authors')
        status = data.get('status')
        submission_date = data.get('submission_date')
        expiration_date = data.get('expiration_date')
        
        cursor = db.cursor()
        

        cursor.execute("""
            UPDATE documents 
            SET title = %s
            WHERE id = %s
        """, (title, document_id))
        

        cursor.execute("""
            UPDATE document_metadata 
            SET category = %s, authors = %s, status = %s, 
                submission_date = %s, expiration_date = %s
            WHERE document_id = %s
        """, (category, authors, status, submission_date, expiration_date, document_id))
        
        db.commit()
        

        log_activity(session['user_id'], f"Edited document ID: {document_id}")
        
        return jsonify({"success": True})
    
    except Error as e:
        if db:
            db.rollback()
        print(f"Error updating document: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
    
    finally:
        if db.is_connected():
            cursor.close()
        if db.is_connected():
            db.close()

@app.route('/view-document/<int:document_id>', methods=['GET'])
def view_document(document_id):
    if 'user_id' not in session: 
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('login'))
    
    cursor = db.cursor(dictionary=True)
    user_id = session['user_id']

    try:

        cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()

        full_name = user['full_name'] if user else 'User'


        cursor.execute("SELECT * FROM user_profiles WHERE user_id = %s", (user_id,))
        profile = cursor.fetchone()

        if not profile:
            profile = {
                'user_id': user_id,
                'first_name': user.get('full_name', '').split(' ')[0] if user.get('full_name') else '',
                'last_name': ' '.join(user.get('full_name', '').split(' ')[1:]) if user.get('full_name') else '',
                'email': user.get('email', ''),
                'phone': '',
                'position': 'Coordinator',
                'institution': '',
                'office_address': '',
                'profile_picture': 'avatar.png'
            }

            insert_fields = ', '.join(profile.keys())
            insert_values = ', '.join(['%s'] * len(profile))
            cursor.execute(
                f"INSERT INTO user_profiles ({insert_fields}) VALUES ({insert_values})",
                tuple(profile.values())
            )
            db.commit()

        profile['profile_picture_url'] = url_for('static', filename=f'uploads/profile_pictures/{profile["profile_picture"]}') if profile.get('profile_picture') else url_for('static', filename='avatar.png')


        cursor.execute("""
            SELECT d.id, d.title, d.type, 
                dm.category, dm.authors, dm.status, 
                dm.submission_date, dm.expiration_date,
                dm.file_path
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            WHERE d.id = %s
        """, (document_id,))

        document = cursor.fetchone()  # First fetch the document


        if document and document.get('file_path'):
            app.logger.info(f"Document ID: {document_id}")
            app.logger.info(f"File path from DB: {document['file_path']}")
            app.logger.info(f"File extension check: {document['file_path'].lower().endswith('.pdf')}")
            expected_path = os.path.join(app.static_folder, 'documents', document['file_path'].split('/')[-1])
            app.logger.info(f"Full expected path: {expected_path}")
            app.logger.info(f"File exists: {os.path.exists(expected_path)}")

        if not document:
            flash('Document not found', 'error')
            return redirect(url_for('tables'))
        
        log_activity(user_id, f"Viewed document: {document['title']} (ID: {document_id})")

        return render_template('view_document.html', document=document, full_name=full_name, profile=profile)
    
    except Error as e:
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('tables'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

        
@app.route('/edit-document/<int:document_id>', methods=['GET', 'POST'])
def edit_document(document_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('tables'))
    
    try:
        cursor = db.cursor(dictionary=True)
        
        if request.method == 'POST':

            title = request.form.get('title')
            category = request.form.get('category')
            authors = request.form.get('authors')
            status = request.form.get('status')
            submission_date = request.form.get('submission_date')
            expiration_date = request.form.get('expiration_date')
            doc_type = request.form.get('type')
            

            cursor.execute("""
                UPDATE documents 
                SET title = %s, type = %s
                WHERE id = %s
            """, (title, doc_type, document_id))
            

            cursor.execute("""
                UPDATE document_metadata 
                SET category = %s, authors = %s, status = %s, 
                    submission_date = %s, expiration_date = %s
                WHERE document_id = %s
            """, (category, authors, status, submission_date, expiration_date, document_id))
            
            db.commit()
            
            log_activity(session['user_id'], f"Edited document ID: {document_id}")
            
            flash('Document updated successfully', 'success')
            return redirect(url_for('tables'))
        

        cursor.execute("""
            SELECT d.id, d.title, d.type, 
                   dm.category, dm.authors, dm.status, 
                   dm.submission_date, dm.expiration_date
            FROM documents d
            JOIN document_metadata dm ON d.id = dm.document_id
            WHERE d.id = %s
        """, (document_id,))
        
        document = cursor.fetchone()
        
        if not document:
            flash('Document not found', 'error')
            return redirect(url_for('tables'))
        

        if document['submission_date']:
            document['submission_date'] = document['submission_date'].strftime('%Y-%m-%d')
        if document['expiration_date']:
            document['expiration_date'] = document['expiration_date'].strftime('%Y-%m-%d')
        

        user_id = session['user_id']
        cursor.execute("SELECT full_name FROM users WHERE id = %s", (user_id,))
        user = cursor.fetchone()
        full_name = user['full_name'] if user else 'User'
        
        return render_template('edit_document.html', document=document, full_name=full_name)
    
    except Error as e:
        flash(f'Error: {str(e)}', 'error')
        return redirect(url_for('tables'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()
            
@app.route('/approve-document/<int:document_id>', methods=['POST'])
def approve_document(document_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('tables'))
    
    try:
        cursor = db.cursor()

        cursor.execute("""
            UPDATE document_metadata
            SET status = 'Approved'
            WHERE document_id = %s
        """, (document_id, ))

        db.commit()

        log_activity(session['user_id'], f"Approved document ID: {document_id}")

        flash('Document approved successfully', 'success')
        return redirect(url_for('tables'))
    
    except Error as e:
        flash(f'Error approving document: {str(e)}', 'error')
        return redirect(url_for('tables'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

@app.route('/reject-document/<int:document_id>', methods=['POST'])
def reject_document(document_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('tables'))
    
    try:
        cursor = db.cursor()
        
        cursor.execute("""
            UPDATE document_metadata 
            SET status = 'Rejected'
            WHERE document_id = %s
        """, (document_id,))
        
        db.commit()
        
        log_activity(session['user_id'], f"Rejected document ID: {document_id}")
        
        flash('Document rejected successfully', 'success')
        return redirect(url_for('tables'))
    
    except Error as e:
        flash(f'Error rejecting document: {str(e)}', 'error')
        return redirect(url_for('tables'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

@app.route('/delete-document/<int:document_id>', methods=['POST'])
def delete_document(document_id):
    is_ajax = request.headers.get('X-Requested-With') == 'XMLHttpRequest'
    
    if 'user_id' not in session:
        if is_ajax:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        error_msg = 'Database connection failed'
        if is_ajax:
            return jsonify({'success': False, 'error': error_msg}), 500
        flash(error_msg, 'error')
        return redirect(url_for('tables'))
    
    try:
        cursor = db.cursor()
        
        cursor.execute("SELECT id FROM documents WHERE id = %s", (document_id,))
        if not cursor.fetchone():
            error_msg = 'Document not found'
            if is_ajax:
                return jsonify({'success': False, 'error': error_msg}), 404
            flash(error_msg, 'error')
            return redirect(url_for('tables'))
        
        log_activity(session['user_id'], f"Deleted document ID: {document_id}")
        
        cursor.execute("DELETE FROM document_metadata WHERE document_id = %s", (document_id,))
        cursor.execute("DELETE FROM documents WHERE id = %s", (document_id,))
        db.commit()
        
        if is_ajax:
            return jsonify({
                'success': True,
                'message': 'Document deleted successfully',
                'document_id': document_id
            })
        
        flash('Document deleted successfully', 'success')
        return redirect(url_for('tables'))
    
    except Error as e:
        db.rollback()
        error_msg = f'Error deleting document: {str(e)}'
        print(f"Delete error: {error_msg}")
        
        if is_ajax:
            return jsonify({'success': False, 'error': error_msg}), 500
        
        flash(error_msg, 'error')
        return redirect(url_for('tables'))
    
    finally:
        if db and db.is_connected():
            cursor.close()
            db.close()

@app.route('/activity-logs')
def activity_logs():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    db = get_db_connection()
    if not db:
        flash('Database Connection Failed', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        cursor = db.cursor(dictionary=True)
        

        cursor.execute("""
            SELECT al.id, al.timestamp, al.activity, u.full_name as user_name
            FROM activity_logs al
            JOIN users u ON al.user_id = u.id
            ORDER BY al.timestamp DESC
            LIMIT 100
        """)
        
        logs = cursor.fetchall()
        
        return render_template('activity_logs.html', logs=logs)
    
    except Error as e:
        flash(f'Error retrieving activity logs: {str(e)}', 'error')
        return redirect(url_for('dashboard'))
    
    finally:
        if db.is_connected():
            cursor.close()
            db.close()

if __name__ == '__main__':
    print("Registered routes:")
    for rule in app.url_map.iter_rules():
        print(f"  {rule}")
    app.run(debug=True)